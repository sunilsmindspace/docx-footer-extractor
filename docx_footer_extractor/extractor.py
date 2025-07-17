"""
DOCX Footer Extractor

A Python library for extracting metadata from DOCX file footers using parallel processing.
Supports Python 3.9+
"""

import os
from docx import Document
from concurrent.futures import ThreadPoolExecutor
from typing import List, Dict, Union
import threading

__version__ = "1.0.0"
__author__ = "Sunil K Sundaram"
__email__ = "sunilsmindspace@gmail.com"

class DocxFooterExtractor:
    """
    A class for extracting metadata from DOCX file footers using parallel processing.
    
    This extractor processes DOCX files to find key-value pairs in footer text and tables.
    It supports both folder-based processing and specific file list processing.
    """
    
    def __init__(self, max_workers: int = None):
        """
        Initialize the DOCX footer metadata extractor.
        
        Args:
            max_workers (int, optional): Maximum number of worker threads. 
                                       If None, uses default ThreadPoolExecutor behavior.
        """
        self.max_workers = max_workers
        self.lock = threading.Lock()
    
    def _extract_key_values(self, text: str) -> Dict[str, str]:
        """
        Extract key-value pairs from text where pairs are separated by colons.
        
        Args:
            text (str): Input text containing key-value pairs
            
        Returns:
            Dict[str, str]: Dictionary of key-value pairs
        """
        result = {}
        for line in text.split('\n'):
            if ':' in line:
                key, value = line.split(':', 1)
                result[key.strip()] = value.strip()
        return result
    
    def _get_kv(self, data: List[tuple]) -> Dict[str, str]:
        """
        Extract key-value pairs from a list of tuples.
        
        Args:
            data (List[tuple]): List of tuples containing text data
            
        Returns:
            Dict[str, str]: Dictionary containing all extracted key-value pairs
        """
        doc_info = {}
        
        for entry in data:
            for part in entry:
                if part and part.strip():  # skip empty strings
                    doc_info.update(self._extract_key_values(part))
        
        return doc_info
    
    def _process_single_file(self, filepath: str) -> Dict[str, Union[str, Dict[str, str]]]:
        """
        Process a single DOCX file and extract metadata from footers.
        
        Args:
            filepath (str): Full path to the DOCX file to process
            
        Returns:
            Dict[str, Union[str, Dict[str, str]]]: Dictionary containing filename and metadata
        """
        filename = os.path.basename(filepath)
        
        try:
            doc = Document(filepath)
            all_metadata = {}
            
            # Get all sections in the document
            for section in doc.sections:
                footer = section.footer
                
                # Extract text from footer paragraphs
                footer_text = []
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():  # Only add non-empty paragraphs
                        footer_text.append(paragraph.text.strip())
                
                # Process footer text for key-value pairs
                for text in footer_text:
                    all_metadata.update(self._extract_key_values(text))
                
                # Extract text from footer tables
                for table in footer.tables:
                    table_data = [tuple(c.text for c in r.cells) for r in table.rows]
                    table_metadata = self._get_kv(table_data)
                    all_metadata.update(table_metadata)
            
            return {
                "filename": filename,
                "metadata": all_metadata
            }
            
        except Exception as e:
            # Return filename with error information
            return {
                "filename": filename,
                "metadata": {"error": str(e)}
            }
    
    def _get_file_list(self, input_source: Union[str, List[str]]) -> List[str]:
        """
        Get list of DOCX files from input source.
        
        Args:
            input_source (Union[str, List[str]]): Either a folder path or list of file paths
            
        Returns:
            List[str]: List of full file paths
            
        Raises:
            FileNotFoundError: If folder doesn't exist
            NotADirectoryError: If path is not a directory
            ValueError: If no valid files found
            TypeError: If input type is invalid
        """
        if isinstance(input_source, str):
            # Input is a folder path
            if not os.path.exists(input_source):
                raise FileNotFoundError(f"Folder not found: {input_source}")
            
            if not os.path.isdir(input_source):
                raise NotADirectoryError(f"Path is not a directory: {input_source}")
            
            # Get all .docx files in the folder
            docx_files = []
            for filename in os.listdir(input_source):
                if filename.endswith('.docx'):
                    docx_files.append(os.path.join(input_source, filename))
            
            if not docx_files:
                raise ValueError(f"No .docx files found in folder: {input_source}")
            
            return docx_files
        
        elif isinstance(input_source, list):
            # Input is a list of file paths
            if not input_source:
                raise ValueError("File list is empty")
            
            # Validate files exist and are .docx files
            valid_files = []
            for filepath in input_source:
                if not os.path.exists(filepath):
                    print(f"Warning: File not found: {filepath}")
                    continue
                
                if not filepath.endswith('.docx'):
                    print(f"Warning: Not a .docx file: {filepath}")
                    continue
                
                valid_files.append(filepath)
            
            if not valid_files:
                raise ValueError("No valid .docx files found in the provided list")
            
            return valid_files
        
        else:
            raise TypeError("Input must be either a folder path (string) or list of file paths")
    
    def extract(self, input_source: Union[str, List[str]], verbose: bool = True) -> List[Dict[str, Union[str, Dict[str, str]]]]:
        """
        Extract metadata from DOCX files using parallel processing.
        
        Args:
            input_source (Union[str, List[str]]): Either a folder path (string) or list of file paths (list)
            verbose (bool): Whether to print processing information
            
        Returns:
            List[Dict[str, Union[str, Dict[str, str]]]]: List of dictionaries, each containing 'filename' and 'metadata' keys
            
        Raises:
            FileNotFoundError: If folder doesn't exist
            NotADirectoryError: If path is not a directory
            ValueError: If no valid files found
            TypeError: If input type is invalid
        """
        # Get list of files to process
        file_list = self._get_file_list(input_source)
        
        if verbose:
            print(f"Found {len(file_list)} .docx files to process:")
            for filepath in file_list:
                print(f"  - {os.path.basename(filepath)}")
        
        # Process files in parallel
        results = []
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tasks
            future_to_file = {
                executor.submit(self._process_single_file, filepath): filepath 
                for filepath in file_list
            }
            
            # Collect results as they complete
            for future in future_to_file:
                try:
                    result = future.result()
                    results.append(result)
                    
                    # Thread-safe printing
                    if verbose:
                        with self.lock:
                            print(f"Processed: {result['filename']}")
                            if result['metadata'] and 'error' not in result['metadata']:
                                print(f"  Found {len(result['metadata'])} metadata entries")
                            elif 'error' in result['metadata']:
                                print(f"  Error: {result['metadata']['error']}")
                            else:
                                print(f"  No metadata found")
                                
                except Exception as e:
                    filepath = future_to_file[future]
                    filename = os.path.basename(filepath)
                    if verbose:
                        with self.lock:
                            print(f"Error processing {filename}: {e}")
                    results.append({
                        "filename": filename,
                        "metadata": {"error": str(e)}
                    })
        
        return results
    
    def save(self, results: List[Dict[str, Union[str, Dict[str, str]]]], output_file: str = "metadata_results.txt") -> None:
        """
        Save the extraction results to a text file.
        
        Args:
            results (List[Dict[str, Union[str, Dict[str, str]]]]): List of dictionaries from extract_metadata
            output_file (str): Name of the output file
        """
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("DOCX Footer Metadata Extraction Results\n")
            f.write("=" * 50 + "\n\n")
            
            for result in results:
                f.write(f"File: {result['filename']}\n")
                f.write("-" * 30 + "\n")
                
                metadata = result['metadata']
                if metadata:
                    for key, value in metadata.items():
                        f.write(f"{key}: {value}\n")
                else:
                    f.write("No metadata found\n")
                
                f.write("\n")
        
        print(f"Results saved to: {output_file}")
